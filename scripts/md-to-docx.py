"""
Convert client-progress-update.md → MoveSmart-Rentals-Progress-Update.docx
Supports: H1–H3, paragraphs, bullets, GFM tables, bold/italic, hr, code spans,
markdown links [text](url), emoji status indicators preserved as text.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = Path(r"e:\MoveSmart-Rentals\docs\client-progress-update.md")
DST = Path(r"e:\MoveSmart-Rentals\docs\MoveSmart-Rentals-Progress-Update.docx")

BRAND_NAVY = RGBColor(0x0B, 0x1D, 0x3A)
BRAND_EMERALD = RGBColor(0x10, 0xB9, 0x81)
BRAND_GOLD = RGBColor(0xC8, 0xA2, 0x4F)
SLATE = RGBColor(0x47, 0x55, 0x69)
MUTED = RGBColor(0x94, 0xA3, 0xB8)

INLINE_RE = re.compile(
    r"(\*\*(?P<bold>.+?)\*\*"
    r"|\*(?P<italic>.+?)\*"
    r"|`(?P<code>.+?)`"
    r"|\[(?P<linktext>.+?)\]\((?P<url>[^)]+)\))"
)


def add_inline(paragraph, text):
    """Parse a single line of inline markdown (bold/italic/code/links) into runs."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        if m.group("bold") is not None:
            run = paragraph.add_run(m.group("bold"))
            run.bold = True
        elif m.group("italic") is not None:
            run = paragraph.add_run(m.group("italic"))
            run.italic = True
        elif m.group("code") is not None:
            run = paragraph.add_run(m.group("code"))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = BRAND_NAVY
        elif m.group("linktext") is not None:
            run = paragraph.add_run(m.group("linktext"))
            run.font.color.rgb = BRAND_EMERALD
            run.underline = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C8A24F")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def build_table(doc, rows):
    """rows is a list of lists; first row is header."""
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""  # clear default empty paragraph text
            p = cell.paragraphs[0]
            add_inline(p, cell_text.strip())
            for run in p.runs:
                run.font.size = Pt(10)
                if i == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if i == 0:
                set_cell_shading(cell, "0B1D3A")
    return table


def style_heading(paragraph, level):
    for run in paragraph.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.size = Pt(26)
            run.bold = True
            run.font.color.rgb = BRAND_NAVY
        elif level == 2:
            run.font.size = Pt(18)
            run.bold = True
            run.font.color.rgb = BRAND_NAVY
        elif level == 3:
            run.font.size = Pt(13)
            run.bold = True
            run.font.color.rgb = BRAND_EMERALD


def main():
    if not SRC.exists():
        print(f"Source not found: {SRC}", file=sys.stderr)
        sys.exit(1)

    text = SRC.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()

    # Default body style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.3

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Blank line
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            add_horizontal_rule(doc)
            i += 1
            continue

        # Headings
        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, stripped[2:])
            style_heading(p, 1)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(12)
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph()
            add_inline(p, stripped[3:])
            style_heading(p, 2)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            i += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph()
            add_inline(p, stripped[4:])
            style_heading(p, 3)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # GFM Table: header row | separator row | data rows
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip()):
            rows = []
            # collect header
            rows.append([c.strip() for c in stripped.strip("|").split("|")])
            j = i + 2  # skip separator
            while j < len(lines) and lines[j].strip().startswith("|"):
                row = [c.strip() for c in lines[j].strip().strip("|").split("|")]
                rows.append(row)
                j += 1
            build_table(doc, rows)
            doc.add_paragraph()  # spacer
            i = j
            continue

        # Bullet list
        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
            for run in p.runs:
                run.font.color.rgb = SLATE
                run.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # Numbered list
        m_num = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m_num:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, m_num.group(2))
            for run in p.runs:
                run.font.color.rgb = SLATE
                run.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # Italic-only line (e.g. "*Prepared by the dev team...*")
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.italic = True
            run.font.color.rgb = MUTED
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            continue

        # Plain paragraph (may span multiple lines until blank)
        para_lines = [stripped]
        j = i + 1
        while j < len(lines) and lines[j].strip() and not lines[j].strip().startswith(
            ("#", "-", "*", "|", ">", "1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.")
        ) and lines[j].strip() != "---":
            para_lines.append(lines[j].strip())
            j += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(para_lines))
        for run in p.runs:
            run.font.color.rgb = SLATE
        i = j

    DST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DST))
    print(f"Wrote: {DST} ({DST.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
